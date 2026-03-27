from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import json


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_code_block(doc: Document, text: str) -> None:
    # Simple monospace-like formatting: keep it as a normal paragraph.
    p = doc.add_paragraph()
    p_run = p.add_run(text)
    p_run.font.name = "Consolas"
    p_run.font.size = Pt(9)


def read_json(path: str) -> dict:
    try:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    except Exception:
        return {}


def build_repo_tree() -> str:
    return "\n".join(
        [
            "agents/",
            "  __init__.py",
            "  data_collection_agent.py",
            "  data_quality_agent.py",
            "  annotation_agent.py",
            "  al_agent.py",
            "  utils.py",
            "app.py",
            "config.yaml",
            "run_pipeline.py",
            "README.md",
            "requirements.txt",
            "",
            "scripts/",
            "  generate_notebooks.py",
            "  generate_system_docx.py",
            "",
            "skills/",
            "  data_collection_agent/SKILL.md",
            "  data_quality_agent/SKILL.md",
            "  annotation_agent/SKILL.md",
            "  al_agent/SKILL.md",
            "",
            "notebooks/",
            "  eda.ipynb",
            "  al_experiment.ipynb",
            "  quality_analysis.ipynb",
            "  annotation_analysis.ipynb",
            "",
            "data/",
            "  raw/   (партиции/сборка сырья)",
            "  labeled/ (промежуточные/финальные таблицы)",
            "",
            "models/",
            "  final_model.pkl",
            "  vectorizer.pkl",
            "",
            "reports/",
            "  quality_report.md",
            "  annotation_report.md",
            "  al_report.md",
            "  final_report.md",
            "  classification_report.csv/.json",
            "  confusion_matrix.csv",
            "",
            "artifacts/",
            "  run_<run_id>/ (run-time артефакты)",
            "",
            "review_queue.csv",
            "review_queue_corrected.csv",
        ]
    )


def build_pipeline_flow() -> str:
    return "\n".join(
        [
            "1) Collection",
            "   DataCollectionAgent.run(sources) -> raw_df",
            "",
            "2) Cleaning / Quality",
            "   DataQualityAgent.detect_issues(raw_df) -> dq_report",
            "   DataQualityAgent.fix(raw_df, strategy) -> clean_df",
            "   DataQualityAgent.compare(raw_df, clean_df) -> compare_table",
            "",
            "3) Annotation + HITL queue",
            "   AnnotationAgent.auto_label(clean_df) -> labeled_df (label_auto, confidence)",
            "   AnnotationAgent.generate_spec(labeled_df, task) -> annotation_spec.md",
            "   AnnotationAgent.build_review_queue(labeled_df, threshold, prioritize_rare) -> review_queue.csv",
            "   Human edits review_queue_corrected.csv (label_human)",
            "   Resume: pipeline applies corrected labels -> label_final",
            "",
            "4) Active Learning (Track A)",
            "   prepare_al_dataset():",
            "     - dedup_before_split (by text) to avoid leakage",
            "     - unknown_policy: drop/cap/keep",
            "     - min_class_count filter for rare classes",
            "   ActiveLearningAgent.run_cycle() with strategies entropy & random",
            "",
            "5) Training",
            "   Train final TF-IDF + LogisticRegression model",
            "   Save: models/final_model.pkl + models/vectorizer.pkl",
            "",
            "6) Reporting",
            "   Top-level Markdown reports in reports/",
            "   per-class diagnostics (classification_report + confusion_matrix)",
        ]
    )


def main() -> None:
    doc = Document()
    add_title(doc, "Genome AI Agents — Системная документация (мультиагентный пайплайн)")
    add_par(doc, f"Сгенерировано: {datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph()

    add_heading(doc, "1. Цели проекта", level=1)
    add_par(
        doc,
        "Цель проекта — построить воспроизводимую мультиагентную систему для подготовки данных и обучения модели "
        "классификации, с точкой контроля качества «человек-в-контуре» (HITL). Система автоматизирует: сбор данных "
        "из нескольких источников, детекцию и исправление проблем качества, авто-разметку, формирование очереди на "
        "ручную проверку (HITL), Active Learning цикл и финальное обучение модели с сохранением артефактов и отчётов.",
    )

    add_heading(doc, "2. Общая архитектура (модули и агенты)", level=1)
    add_heading(doc, "2.1. Оркестрация", level=2)
    add_par(
        doc,
        "Оркестрация реализована в `run_pipeline.py`. Он вызывает агентов последовательно, "
        "сериализует результаты в CSV/JSON и генерирует отчеты. "
        "Команда запуска: `py run_pipeline.py`. Поддерживается resume для HITL: "
        "`py run_pipeline.py --resume-hitl`.",
    )

    add_heading(doc, "2.2. Агентный состав", level=2)
    add_par(doc, "- `DataCollectionAgent` — сбор и унификация датасета.")
    add_par(doc, "- `DataQualityAgent` — детект и исправление проблем качества данных.")
    add_par(doc, "- `AnnotationAgent` — авто-разметка, генерация спецификации и экспорт в LabelStudio, формирование HITL очереди.")
    add_par(doc, "- `ActiveLearningAgent` — Active Learning (Track A), обучение базовой модели и выбор информативных примеров.")

    add_heading(doc, "2.3. Skill-инфраструктура (для Claude Code)", level=2)
    add_par(
        doc,
        "В папке `skills/` находятся `SKILL.md` документы, описывающие контракт методов каждого агента. "
        "Эти файлы предназначены для внешней агентной/IDE интеграции и не обязательны для работы Python-пайплайна, "
        "но обеспечивают строгую документированность API агентов.",
    )

    add_heading(doc, "3. Репозиторий: полная структура", level=1)
    add_par(doc, "Текущая структура (инвариантная часть):")
    add_code_block(doc, build_repo_tree())

    add_heading(doc, "4. Контракты данных (unified schema)", level=1)
    add_par(
        doc,
        "Агенты 1–2 и последующие шаги используют единый контракт таблицы pandas.DataFrame. "
        "Фиксированные колонки:\n"
        "- `text`: текстовое представление примера.\n"
        "- `audio`: поле для аудио (в current template текстовый пайплайн; обычно null).\n"
        "- `image`: поле для изображений (обычно null).\n"
        "- `label`: исходная метка (может быть `unknown` до разметки).\n"
        "- `source`: происхождение (`hf:...`, `kaggle:...`, `scrape:...`).\n"
        "- `collected_at`: UTC timestamp в ISO-формате.",
    )

    add_heading(doc, "5. Пайплайн: пошаговое описание работы", level=1)
    add_code_block(doc, build_pipeline_flow())

    add_heading(doc, "6. История улучшений (что делали агенты и как рос результат)", level=1)
    add_par(
        doc,
        "Этот раздел описывает последовательность работ «как было» → «что исправили» → «как стало». "
        "Цель — зафиксировать причинно‑следственную связь между изменениями в пайплайне и качеством модели.",
    )

    add_heading(doc, "6.1. Этап 1 — Сбор данных (DataCollectionAgent)", level=2)
    add_par(
        doc,
        "Агент `DataCollectionAgent` собирает данные из источников (HuggingFace/Kaggle/скрейпинг) и приводит их "
        "к единой схеме колонок (`text`, `label`, `source`, `collected_at`, плюс `audio/image` как null для текстового шаблона).",
    )
    add_par(doc, "Артефакты: `data/raw/data_raw.csv`, `artifacts/run_<run_id>/data_raw.csv`.")

    add_heading(doc, "6.2. Этап 2 — Контроль качества (DataQualityAgent)", level=2)
    add_par(
        doc,
        "Агент `DataQualityAgent` строит quality‑отчёт и применяет стратегии очистки: пропуски, дубликаты, выбросы. "
        "В исходных данных обнаруживались значимые дубликаты по `text` и сильный дисбаланс, включая большой класс `unknown`.",
    )
    add_par(doc, "Артефакты: `reports/quality_report.md`, `artifacts/run_<run_id>/reports/quality_report.json`.")

    add_heading(doc, "6.3. Этап 3 — Авторазметка + HITL (AnnotationAgent)", level=2)
    add_par(
        doc,
        "`AnnotationAgent` присваивает предварительные метки (`label_auto`) и оценку уверенности (`confidence`). "
        "Далее строится очередь на ручную проверку (HITL) по low‑confidence примерам. "
        "После внесения правок человеком пайплайн применяет `label_human` и формирует `label_final`.",
    )
    add_par(
        doc,
        "Артефакты: `review_queue.csv`, `review_queue_corrected.csv`, `artifacts/run_<run_id>/data_labeled_final.csv`.",
    )

    add_heading(doc, "6.4. Этап 4 — Подготовка датасета для обучения (leakage‑safe prep)", level=2)
    add_par(
        doc,
        "Перед обучением применяется подготовка для исключения утечек и снижения шума:\n"
        "- дедупликация по `text` до split (чтобы одинаковые строки не попадали в train и test)\n"
        "- политика по `unknown` (в текущем конфиге: `drop` для обучения)\n"
        "- фильтр слишком редких классов `min_class_count`",
    )

    add_heading(doc, "6.5. Проблема — коллапс предсказаний (модель предсказывала один класс)", level=2)
    add_par(
        doc,
        "На ранней версии baseline‑модели наблюдался симптом «prediction collapse»: "
        "модель предсказывала практически один класс для всех объектов теста. "
        "Это фиксировалось в per‑class отчёте и подтверждалось диагностикой.",
    )

    add_heading(doc, "6.6. Исправление — смена представления признаков (char‑ngrams вместо word‑ngrams)", level=2)
    add_par(
        doc,
        "Было выявлено, что входные тексты имеют структурный вид (много повторяющихся ключей и значений типа `feature=0/1`). "
        "Word‑TF‑IDF почти не создавал полезного разделяющего сигнала, поэтому логистическая регрессия сваливалась "
        "в «безопасный» один класс. Решение: перейти на character n‑grams (`char_wb`, 3–5) и solver `saga` "
        "с увеличенным `max_iter` и настройкой `C`.",
    )
    add_par(
        doc,
        "Изменения внесены в: `agents/al_agent.py` (финальная модель) и `agents/annotation_agent.py` (авторазметка).",
    )

    add_heading(doc, "6.7. Результат — метрики после исправления", level=2)
    diag = read_json("reports/model_diagnostics.json")
    if diag:
        acc = diag.get("accuracy")
        f1m = diag.get("f1_macro")
        collapse = diag.get("prediction_collapse")
        add_par(
            doc,
            f"Снимок метрик (из `reports/model_diagnostics.json`): accuracy={acc}, f1_macro={f1m}, prediction_collapse={collapse}.",
        )
        add_par(
            doc,
            "Дополнительно сохраняются распределения классов train/test и распределение предсказаний модели, "
            "чтобы быстро диагностировать деградации при изменении данных/конфига.",
        )
    else:
        add_par(
            doc,
            "Снимок метрик берётся из `reports/model_diagnostics.json` (если файл доступен).",
        )

    add_heading(doc, "6.8. Unsupervised ML для `unknown` — ускорение разметки (HITL)", level=2)
    add_par(
        doc,
        "Для большого класса `unknown` добавлены инструменты кластеризации и визуализации:\n"
        "- `scripts/cluster_unknowns.py` → `reports/unknown_clusters.csv` и `reports/unknown_clusters.md`\n"
        "- `scripts/visualize_unknown_clusters.py` → `reports/unknown_clusters_plot.png`\n"
        "Это позволяет эксперту просматривать группы похожих текстов и присваивать метки пакетно.",
    )

    add_heading(doc, "7. Диагностика коллапса предсказаний и метрики (артефакт)", level=1)
    add_par(
        doc,
        "Система сохраняет машинно-читаемую диагностику модели в `reports/model_diagnostics.json` (и в run-папке "
        "`artifacts/run_<run_id>/reports/model_diagnostics.json`). В файле фиксируются:\n"
        "- распределение классов в train/test\n"
        "- распределение предсказаний модели\n"
        "- список уникальных предсказанных меток\n"
        "- флаг `prediction_collapse` (true, если модель предсказывает 1 класс или меньше)\n"
        "Это позволяет быстро обнаружить ситуацию «модель предсказывает один класс» до интерпретации "
        "classification_report/confusion_matrix.",
    )

    add_heading(doc, "8. Линия данных: куда пишутся файлы и что они означают", level=1)
    add_par(
        doc,
        "Пайплайн пишет артефакты в двух местах:\n"
        "- `artifacts/run_<run_id>/` — полностью воспроизводимый слепок конкретного прогона (включая отчёты)\n"
        "- «latest» копии в корне проекта (`data/`, `reports/`, `models/`, `review_queue.csv`) — удобны для просмотра\n"
        "Ниже — перечень файлов и их смысл.",
    )

    add_heading(doc, "8.1. Папка `data/raw/` (сырьё)", level=2)
    add_par(doc, "`data/raw/data_raw.csv` — объединённые данные после `DataCollectionAgent`, до очистки.")
    add_par(
        doc,
        "Колонки (единый контракт): `text`, `audio`, `image`, `label`, `source`, `collected_at`.\n"
        "- `text`: исходный текст/описание (для табличных источников — строка вида `key=value; ...`).\n"
        "- `label`: исходная метка из источника (может быть `unknown`).\n"
        "- `source`: откуда пришла строка (hf/kaggle/scrape).\n"
        "- `collected_at`: время сбора (UTC ISO).",
    )

    add_heading(doc, "8.2. Папка `data/labeled/` (очищенные и размеченные таблицы)", level=2)
    add_par(doc, "`data/labeled/data_clean.csv` — результат `DataQualityAgent.fix()` (очистка пропусков/дубликатов/выбросов).")
    add_par(
        doc,
        "`data/labeled/data_auto_labeled.csv` — после `AnnotationAgent.auto_label()`.\n"
        "Добавляются поля:\n"
        "- `label_auto`: предложенная моделью метка\n"
        "- `confidence`: уверенность предложения (0..1)",
    )
    add_par(
        doc,
        "`artifacts/run_<run_id>/data_labeled_final.csv` — финальная таблица после применения HITL (если был `--resume-hitl`).\n"
        "Ключевое поле:\n"
        "- `label_final`: итоговая метка (либо `label_auto`, либо `label_human` где есть правка; low-confidence может становиться `unknown`).",
    )

    add_heading(doc, "8.3. HITL файлы (ручная корректировка)", level=2)
    add_par(
        doc,
        "`review_queue.csv` (и копия в `artifacts/run_<run_id>/review_queue.csv`) — очередь на проверку.\n"
        "Типичные колонки:\n"
        "- `text`: пример\n"
        "- `label_suggested`: предложенная метка\n"
        "- `confidence`: уверенность предложения\n"
        "- `label_human`: пустая колонка для заполнения экспертом",
    )
    add_par(
        doc,
        "`review_queue_corrected.csv` — файл после ручной разметки (обязательная колонка `label_human`).\n"
        "При запуске `py run_pipeline.py --resume-hitl` правки применяются к `label_final`.",
    )

    add_heading(doc, "8.4. Папка `reports/` (отчёты и диагностика)", level=2)
    add_par(
        doc,
        "Ключевые отчёты:\n"
        "- `reports/quality_report.md` — отчёт качества (пропуски/дубликаты/выбросы/дисбаланс)\n"
        "- `reports/annotation_report.md` — отчёт по авторазметке + метрики (в т.ч. kappa после HITL, если есть)\n"
        "- `reports/al_report.md` + `reports/learning_curve.png` — кривая Active Learning (macro‑F1 vs n_labeled)\n"
        "- `reports/classification_report.json/.csv` — per-class precision/recall/F1/support\n"
        "- `reports/confusion_matrix.csv` — матрица ошибок\n"
        "- `reports/model_diagnostics.json` — распределения train/test/pred и флаг `prediction_collapse`\n"
        "- `reports/final_report.md` — краткий итог с метриками",
    )
    add_par(
        doc,
        "Для каждого run существует аналогичная папка `artifacts/run_<run_id>/reports/` с привязкой к конкретному прогону.",
    )

    add_heading(doc, "8.5. Папка `models/` (модель и векторизатор)", level=2)
    add_par(
        doc,
        "`models/final_model.pkl` — сериализованный sklearn Pipeline (TF‑IDF + LogisticRegression).\n"
        "`models/vectorizer.pkl` — сохранённый TF‑IDF шаг (если доступен) для повторного использования в инференсе.",
    )

    add_heading(doc, "8.6. Папка `artifacts/run_<run_id>/` (слепок прогона)", level=2)
    add_par(
        doc,
        "Внутри `artifacts/run_<run_id>/` лежат все промежуточные таблицы и отчёты конкретного запуска:\n"
        "- `data_raw.csv`, `data_clean.csv`, `data_auto_labeled.csv`, `data_labeled_final.csv`\n"
        "- `review_queue.csv`, `review_queue_corrected.csv` (если был resume)\n"
        "- `labelstudio_import.json` — импорт для LabelStudio\n"
        "- `reports/` — отчёты, метрики, матрица ошибок, `model_diagnostics.json` и т.д.",
    )

    add_heading(doc, "9. Unsupervised ML для `unknown` + HITL (кластеризация)", level=1)
    add_par(
        doc,
        "Для работы с большим классом `unknown` добавлен отдельный шаг кластеризации текстов без подтверждённой метки. "
        "Скрипт `scripts/cluster_unknowns.py` группирует `unknown`-строки, а результаты сохраняются в:\n"
        "- `reports/unknown_clusters.csv` (cluster_id, text) — удобно для массовой разметки\n"
        "- `reports/unknown_clusters.md` — краткое описание кластеров: топ-термы и примеры\n"
        "Цель: ускорить разметку (HITL) — эксперт может присваивать диагноз целому кластеру, если он однородный.",
    )
    add_par(doc, "Команда: `py scripts/cluster_unknowns.py --input artifacts/run_<run_id>/data_labeled_final.csv --out-dir reports`")

    add_heading(doc, "10. Визуализация кластеров `unknown`", level=1)
    add_par(
        doc,
        "Для визуального контроля качества кластеризации добавлен скрипт `scripts/visualize_unknown_clusters.py`, "
        "который строит 2D-проекцию TF-IDF признаков методом TruncatedSVD и раскрашивает точки по `cluster_id`.\n"
        "Результат сохраняется в PNG: `reports/unknown_clusters_plot.png`.",
    )
    add_par(doc, "Команда: `py scripts/visualize_unknown_clusters.py --input reports/unknown_clusters.csv --out reports/unknown_clusters_plot.png`")

    add_heading(doc, "11. Детальная спецификация агентов", level=1)

    add_heading(doc, "6.1. DataCollectionAgent", level=2)
    add_par(doc, "Основные методы:")
    add_par(doc, "- `scrape(url, selector, max_items)` -> DataFrame")
    add_par(doc, "- `fetch_api(endpoint, params)` -> DataFrame (best-effort normalizing JSON)")
    add_par(doc, "- `load_dataset(name, split, sample, text_col, label_col, feature_cols, text_prefix)` -> DataFrame")
    add_par(doc, "- `merge(sources)` -> DataFrame (concat + canonicalize)")
    add_par(doc, "Ключевые гарантии:")
    add_par(doc, "- унифицирует схему (включая `audio`/`image` как null для текстового шаблона)")
    add_par(doc, "- сериализует `text/label/source/collected_at` как string dtype")

    add_heading(doc, "6.2. DataQualityAgent", level=2)
    add_par(doc, "Основные методы:")
    add_par(doc, "- `detect_issues(df)` -> dict")
    add_par(doc, "- `fix(df, strategy)` -> DataFrame")
    add_par(doc, "- `compare(df_before, df_after)` -> DataFrame")
    add_par(
        doc,
        "Логика detect/fix:\n"
        "- missing: count/pct по колонкам (для audio/image учитывается как not applicable, если весь столбец null)\n"
        "- duplicates: подсчет дубликатов по `text`\n"
        "- outliers: IQR-based границы по `text_length` (и примеры индексов)\n"
        "- imbalance: распределение по `label` и max/min ratio\n"
        "- стратегии fix:\n"
        "  - missing: `drop` или `fill_unknown`\n"
        "  - duplicates: `drop` или `keep_first`\n"
        "  - outliers: `clip_iqr`, `remove_iqr`, `none`",
    )

    add_heading(doc, "6.3. AnnotationAgent", level=2)
    add_par(doc, "Основные методы:")
    add_par(doc, "- `auto_label(df, modality='text')` -> DataFrame")
    add_par(doc, "- `generate_spec(df, task)` -> str (Markdown)")
    add_par(doc, "- `check_quality(df_labeled)` -> dict (label distribution, confidence, kappa если есть label_human)")
    add_par(doc, "- `export_to_labelstudio(df, out_path)` -> Path (JSON импорт в LabelStudio)")
    add_par(
        doc,
        "HITL очередь:\n"
        "- `build_review_queue(df_labeled, threshold, prioritize_rare=True)` возвращает low-confidence примеры.\n"
        "- Если `prioritize_rare=True`, сортировка сначала по редкости предсказанного класса (`label_frequency`, рассчитанной по всему датасету), затем по возрастанию `confidence`.\n"
        "- Запись в очередь включает: `text`, `label_suggested`, `confidence`, `source`, `collected_at`, `label_human` (пустая).",
    )

    add_heading(doc, "6.4. ActiveLearningAgent (Track A)", level=2)
    add_par(doc, "Основные методы:")
    add_par(doc, "- `fit(labeled_df)` -> sklearn Pipeline (TF-IDF + LogisticRegression)")
    add_par(doc, "- `query(model, pool_df, strategy, k)` -> list[int] индексы")
    add_par(doc, "- `evaluate(model, labeled_df, test_df)` -> dict (accuracy, macro-F1)")
    add_par(doc, "- `run_cycle(...)` -> history list[dict]")
    add_par(doc, "- `report(history, out_path)` -> learning-curve plot (F1 vs n_labeled)")
    add_par(
        doc,
        "Стратегии отбора:"
        "\n- `entropy`: выбирает примеры с максимальной энтропией предсказаний"
        "\n- `margin`: выбирает минимальный разрыв между top-1 и top-2 вероятностями"
        "\n- `random`: случайный выбор",
    )

    add_heading(doc, "12. Конфигурация (`config.yaml`)", level=1)
    add_par(
        doc,
        "Файл `config.yaml` задает:\n"
        "- collection.sources: HF/Kaggle/scrape источники\n"
        "- quality.fix.strategy: правила очистки\n"
        "- annotation.modality, annotation.confidence_threshold, annotation.prioritize_rare\n"
        "- active_learning: model, n_initial, n_iterations, batch_size, strategies + параметры prep\n",
    )
    add_heading(doc, "7.1. AL prep (улучшенный контроль дисбаланса и утечек)", level=2)
    add_par(
        doc,
        "В `run_pipeline.py` добавлен preparer для AL:\n"
        "- `dedup_before_split` (по `text`) — чтобы исключить leakage между train/test\n"
        "- `unknown_policy`: drop | cap | keep\n"
        "- `unknown_cap_ratio`: если cap, ограничиваем долю unknown\n"
        "- `min_class_count`: исключение сверх-редких классов из AL цикла\n",
    )

    add_heading(doc, "13. Human-in-the-loop: рабочий цикл", level=1)
    add_par(
        doc,
        "1) Pipeline запускает автo-разметку и формирует `review_queue.csv` (или в `artifacts/run_<run_id>/`).\n"
        "2) Эксперт заполняет `label_human` и сохраняет как `review_queue_corrected.csv`.\n"
        "3) Повторный запуск: `python run_pipeline.py --resume-hitl` применяет исправления к `label_final`.\n"
        "4) После resume пересчитываются метрики аннотаций (к примеру, Cohen's kappa на reviewed subset).",
    )

    add_heading(doc, "14. Артефакты и отчеты (как проверять сдачу)", level=1)
    add_par(
        doc,
        "Pipeline сохраняет run-time артефакты в `artifacts/run_<run_id>/` и latest-артефакты в директории проекта.\n"
        "Ключевые outputs:\n"
        "- data/raw/data_raw.csv\n"
        "- data/labeled/data_clean.csv, data_auto_labeled.csv, data_labeled_final.csv\n"
        "- reports/quality_report.md, annotation_report.md, al_report.md, final_report.md\n"
        "- reports/classification_report.csv/.json, reports/confusion_matrix.csv\n"
        "- models/final_model.pkl (+ vectorizer.pkl при обучении)\n",
    )

    add_heading(doc, "15. Как запустить (инструкция)", level=1)
    add_par(doc, "1) Установка:")
    add_code_block(doc, "py -m pip install -r requirements.txt")
    add_par(doc, "2) Полный запуск:")
    add_code_block(doc, "py run_pipeline.py")
    add_par(doc, "3) HITL resume:")
    add_code_block(doc, "py run_pipeline.py --resume-hitl")

    add_heading(doc, "16. Примечания по качеству и метрикам", level=1)
    add_par(
        doc,
        "Важно: в мультиклассовой задаче с сильным дисбалансом `accuracy` может быть обманчивой, "
        "поэтому система ориентируется на macro-F1. Для диагностики сохраняются per-class отчеты "
        "и confusion matrix. AL может казаться неэффективным на малом n_labeled — до тех пор, "
        "пока модель не формирует слабые границы классов; поэтому seed размер (n_initial) "
        "и обработка unknown/редких классов критичны.",
    )

    out_path = Path("SYSTEM_DOCUMENTATION.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()

